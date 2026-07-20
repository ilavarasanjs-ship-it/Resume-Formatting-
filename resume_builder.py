"""
resume_builder.py  v4
Matches the gold-standard manual formatted resume exactly.

Verified format rules from manual analysis:
- Default font: Times New Roman 11pt
- Name: 16pt Bold Center  (NOT italic)
- Additional fields: each on its OWN 12pt Bold Center paragraph (NOT combined)
- Summary heading: Bold Left
- Summary body: first word Bold, rest plain, Justify
- Section headings: Bold Left (NOT italic) — Summary, Education and Certifications, Skills, Experience
- Sub-section headings: Bold Left (NOT italic) — Research Experience, Teaching Experience, etc.
- Institution/school: Plain Left
- Degree bullets: Bold, List Paragraph
- Certification bullets: Bold, List Paragraph
- Skills: OPTIONAL sub-category labels in Bold, then plain List Paragraph bullets
- Company line: Bold Left with right-tab at 9360 twips for dates
- Course line: Bold Left (NOT italic), prefixed with "Course: "
- Job title: Bold Italic Left
- Bullet points: PLAIN (not bold), List Paragraph
- Spacers between every logical block
- TWO spacers before Experience section
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


def _set_default_font(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    existing = rPr.find(qn('w:rFonts'))
    if existing is None:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'),    'Times New Roman')
        rPr.insert(0, rFonts)


def _spacer(doc, style="Normal"):
    doc.add_paragraph(style=style)


def _para(doc, text, bold=False, italic=False, size_pt=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, style="Normal"):
    """Add a plain paragraph — bold and italic explicitly off unless specified."""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    if not text:
        return p
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    return p


def _section_heading(doc, text, with_border=True):
    """
    Bold Left paragraph with bottom border line underneath.
    Matches gold standard: single, sz=4, space=1, color=auto.
    Applied to: Summary, Education and Certifications, Skills, Experience,
    Other Research Experience, Teaching Experience, Learning Assistant Experience,
    Leadership Experience, Clinical Experience, Professional Development.
    NOT applied to: Research Experience (first sub-section, no border in gold standard).
    """
    p   = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold   = True
    run.italic = False
    if with_border:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'auto')
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def _summary_body(doc, text):
    """Summary body: first word bold, rest plain, justified."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not text:
        return p
    parts = text.split(' ', 1)
    r1 = p.add_run(parts[0])
    r1.bold   = True
    r1.italic = False
    if len(parts) > 1:
        r2 = p.add_run(' ' + parts[1])
        r2.bold   = False
        r2.italic = False
    return p


def _bullet(doc, text, bold=False):
    """List Paragraph bullet. Plain by default."""
    p   = doc.add_paragraph(style="List Paragraph")
    run = p.add_run(text)
    run.bold   = bold
    run.italic = False
    return p


def _company_line(doc, institution, dates):
    """Bold left with right-aligned dates via tab stop."""
    p   = doc.add_paragraph(style="Normal")
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)

    def bold_r(txt):
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = txt
        if txt and (txt.startswith(' ') or txt.endswith(' ')):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        return r

    def tab_r():
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))
        r.append(rPr)
        r.append(OxmlElement('w:tab'))
        return r

    p._p.append(bold_r(institution))
    if dates:
        p._p.append(tab_r())
        p._p.append(bold_r(dates))
    return p


def build_resume_docx(data):
    """
    Build the formatted resume DOCX.

    data keys:
        name                str
        location            str   (additional field 1)
        commute_info        str   (additional field 2)
        availability        str   (additional field 3)
        summary             str
        education           list {institution, degrees:[str]}
        certifications      list of str
        skills              list of str  (flat — no categories)
        experience_sections list {section_heading, jobs:[{institution,dates,course,title,bullets}]}
    """
    doc = Document()
    _set_default_font(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin  = sec.bottom_margin = Inches(1)

    # ── NAME — Bold only, 16pt, Center ──────────────────────────────────────
    _para(doc, data.get("name",""), bold=True, italic=False,
          size_pt=16, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── ADDITIONAL FIELDS — each on its own 12pt Bold Center paragraph ───────
    for field in [data.get("location",""), data.get("commute_info",""), data.get("availability","")]:
        if field and field.strip():
            _para(doc, field.strip(), bold=True, italic=False,
                  size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    _spacer(doc)

    # ── SUMMARY ─────────────────────────────────────────────────────────────
    summary = (data.get("summary") or "").strip()
    if summary:
        _section_heading(doc, "Summary")
        _summary_body(doc, summary)
        _spacer(doc)

    # ── EDUCATION AND CERTIFICATIONS ────────────────────────────────────────
    education = data.get("education", [])
    certs     = data.get("certifications", [])
    if education or certs:
        _section_heading(doc, "Education and Certifications")
        for edu in education:
            inst = (edu.get("institution") or "").strip()
            if inst:
                _para(doc, inst, bold=False, italic=False)
            for deg in edu.get("degrees", []):
                if (deg or "").strip():
                    _bullet(doc, deg.strip(), bold=True)
            _spacer(doc)
        for cert in certs:
            if (cert or "").strip():
                _bullet(doc, cert.strip(), bold=True)
        if certs:
            _spacer(doc)

    # ── SKILLS ──────────────────────────────────────────────────────────────
    skills = data.get("skills", [])
    if skills:
        _section_heading(doc, "Skills")
        # Check if skills have category prefixes like "Laboratory:", "Bioinformatics:"
        # The extractor returns flat items — detect category-prefixed items
        current_category = None
        for skill in skills:
            s = (skill or "").strip()
            if not s:
                continue
            # Detect category labels: ends with ":" and is short (< 30 chars)
            if s.endswith(':') and len(s) < 30:
                _para(doc, s, bold=True, italic=False)
                current_category = s
            else:
                _bullet(doc, s, bold=False)
        _spacer(doc)

    # ── EXPERIENCE ──────────────────────────────────────────────────────────
    exp_sections = data.get("experience_sections", [])
    if exp_sections:
        # Two spacers before Experience (matches gold standard)
        if skills or certs or education:
            _spacer(doc)  # already have one spacer from skills/certs, add second

        _section_heading(doc, "Experience")

        for section in exp_sections:
            heading = (section.get("section_heading") or "").strip()
            jobs    = section.get("jobs", [])

            # Sub-section heading — Bold Left with border
            # Exception: "Research Experience" (directly after Experience) has no border
            if heading and heading.lower() not in ("experience", ""):
                no_border_headings = {"research experience"}
                use_border = heading.lower() not in no_border_headings
                _section_heading(doc, heading, with_border=use_border)

            for job in jobs:
                institution = (job.get("institution") or "").strip()
                dates       = (job.get("dates") or "").strip()
                course      = (job.get("course") or "").strip()
                title       = (job.get("title") or "").strip()
                bullets     = job.get("bullets", [])

                # Company + dates line — Bold, tab-separated
                if institution or dates:
                    _company_line(doc, institution, dates)

                # Course line — Bold Left, NOT italic
                if course:
                    # Ensure "Course: " prefix
                    course_text = course if course.startswith("Course:") else f"Course: {course}"
                    _para(doc, course_text, bold=True, italic=False)

                # Job title — Bold Italic
                if title:
                    p   = doc.add_paragraph(style="Normal")
                    run = p.add_run(title)
                    run.bold   = True
                    run.italic = True

                # Bullets — PLAIN
                for b in bullets:
                    clean = (b or "").strip().lstrip("•·∙-–•").strip()
                    if clean:
                        _bullet(doc, clean, bold=False)

                _spacer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
